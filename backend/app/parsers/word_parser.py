"""
LeakSight V1 — Word Document Parser

Source: docs/PARSING_SPEC.md (Section 5.4 — Word Parser)
       docs/ARCHITECTURE.md (Section 6.6 — Parsing Layer)

Handles .docx files using python-docx.
Accuracy target: ≥ 85%.

Key behaviors:
  - Extract tables from doc.tables for line items
  - Extract header info from paragraphs
  - Multi-table handling (contracts with multiple pricing tables)
  - Strip bold/italic formatting, preserve content only
  - Never crash on malformed files
"""

import re
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document as DocxDocument

from backend.app.core.logging import get_logger
from backend.app.parsers.base_parser import (
    BaseParser,
    DocType,
    DocumentHeader,
    FailureFlag,
    LineItem,
    ParseResult,
)

logger = get_logger(__name__)

# Regex patterns — same as other parsers for consistency
_VENDOR_PATTERNS = re.compile(
    r"(?:vendor|supplier|party|from|bill\s*(?:to|from))"
    r"\s*[:\-]?\s*(.+)",
    re.IGNORECASE,
)
_DOC_NUMBER_PATTERNS = re.compile(
    r"(?:invoice|inv|po|grn|bill|contract|ref)\s*(?:no|number|#|num|ref)\.?\s*[:\-]?\s*(\S+)",
    re.IGNORECASE,
)
_DATE_PATTERNS = re.compile(
    r"(?:date|dated|effective\s*date|valid\s*from)\s*[:\-]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
    re.IGNORECASE,
)
_TOTAL_PATTERNS = re.compile(
    r"(?:grand\s*total|total\s*amount|net\s*amount|total)\s*[:\-]?\s*[₹$€£]?\s*([\d,]+\.?\d*)",
    re.IGNORECASE,
)
_GST_PATTERN = re.compile(r"\b(\d{2}[A-Z]{5}\d{4}[A-Z]\d[A-Z\d]{2})\b")

_CURRENCY_STRIP = re.compile(r"[₹$€£¥,\s]")

# Column mapping patterns
_ITEM_DESC_PATTERNS = re.compile(
    r"(item|desc|particular|product|service|material|name)", re.IGNORECASE
)
_QUANTITY_PATTERNS = re.compile(r"(qty|quantity|quant|nos)", re.IGNORECASE)
_UNIT_PATTERNS = re.compile(r"(unit|uom|measure)", re.IGNORECASE)
_UNIT_PRICE_PATTERNS = re.compile(
    r"(rate|unit.?price|price|per.?unit)", re.IGNORECASE
)
_LINE_TOTAL_PATTERNS = re.compile(
    r"(amount|total|value|line.?total|net.?amount)", re.IGNORECASE
)

# Header keywords for table header detection
_TABLE_HEADER_KEYWORDS = {
    "item", "description", "desc", "particular", "particulars",
    "quantity", "qty", "unit", "uom",
    "price", "rate", "amount", "total", "value",
    "sl", "sr", "no", "sno",
}


def _parse_numeric(value: Any) -> Decimal | None:
    """Parse a value to Decimal, handling currency symbols and commas."""
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    s = _CURRENCY_STRIP.sub("", s)
    if s.startswith("(") and s.endswith(")"):
        s = "-" + s[1:-1]
    try:
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return None


def _parse_date_value(value: str) -> date | None:
    """Try to parse a date from various formats."""
    s = value.strip()
    for fmt in (
        "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y",
        "%Y-%m-%d", "%m/%d/%Y",
        "%d-%b-%Y", "%d %b %Y", "%d %B %Y",
        "%d/%m/%y", "%d-%m-%y",
    ):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def _cell_text(cell: Any) -> str:
    """Extract plain text from a docx table cell, stripping formatting."""
    # python-docx Cell has .text property that auto-strips formatting
    return str(cell.text).strip() if cell and cell.text else ""


def _is_table_header_row(cells: list[str]) -> bool:
    """Check if a list of cell texts looks like a header row."""
    score = sum(
        1 for v in cells
        if v.lower() in _TABLE_HEADER_KEYWORDS
        or any(kw in v.lower() for kw in _TABLE_HEADER_KEYWORDS)
    )
    return score >= 2


def _map_table_columns(headers: list[str]) -> dict[str, int | None]:
    """Map table column headers to schema field indices."""
    mapping: dict[str, int | None] = {
        "item_desc": None,
        "quantity": None,
        "unit": None,
        "unit_price": None,
        "line_total": None,
    }
    for idx, col in enumerate(headers):
        col_lower = col.strip().lower()
        if mapping["item_desc"] is None and _ITEM_DESC_PATTERNS.search(col_lower):
            mapping["item_desc"] = idx
        elif mapping["quantity"] is None and _QUANTITY_PATTERNS.search(col_lower):
            mapping["quantity"] = idx
        elif mapping["unit"] is None and _UNIT_PATTERNS.search(col_lower):
            mapping["unit"] = idx
        elif mapping["unit_price"] is None and _UNIT_PRICE_PATTERNS.search(col_lower):
            mapping["unit_price"] = idx
        elif mapping["line_total"] is None and _LINE_TOTAL_PATTERNS.search(col_lower):
            mapping["line_total"] = idx
    return mapping


class WordParser(BaseParser):
    """Parser for Word (.docx) documents.

    Uses python-docx to extract tables and paragraphs.
    Tables are the primary source of line items.
    Paragraphs are used for header information extraction.
    """

    @property
    def supported_formats(self) -> list[str]:
        return [".docx"]

    @property
    def parser_name(self) -> str:
        return "word_parser_v1"

    def parse(self, file_path: Path, document_id: UUID) -> ParseResult:
        """Parse a Word document into the normalized intermediate schema."""
        failure_flags: list[FailureFlag] = []
        line_items: list[LineItem] = []
        header = DocumentHeader()
        confidence = 1.0
        raw_data: dict = {}

        try:
            doc = DocxDocument(str(file_path))

            # ── Extract header info from paragraphs ───────────────
            paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            raw_data["paragraph_count"] = len(paragraphs_text)
            header = self._extract_header_info(paragraphs_text, failure_flags)

            # ── Extract tables ────────────────────────────────────
            raw_data["table_count"] = len(doc.tables)

            if not doc.tables:
                failure_flags.append(FailureFlag(
                    severity="WARNING",
                    code="NO_TABLES_FOUND",
                    message="No tables found in Word document",
                ))
                confidence -= 0.20

            for table in doc.tables:
                items = self._parse_table(table, len(line_items))
                line_items.extend(items)

            # ── Confidence adjustments ────────────────────────────
            if not line_items and doc.tables:
                failure_flags.append(FailureFlag(
                    severity="ERROR",
                    code="MISSING_LINE_ITEMS",
                    message="Tables found but no line items could be extracted",
                ))
                confidence = max(confidence - 0.40, 0.0)
            elif line_items:
                incomplete_count = sum(
                    1 for li in line_items
                    if li.unit_price is None or li.quantity is None
                )
                deduction = min(incomplete_count * 0.02, 0.30)
                confidence -= deduction
                if incomplete_count > 0:
                    failure_flags.append(FailureFlag(
                        severity="WARNING",
                        code="MISSING_UNIT_PRICE",
                        message=f"{incomplete_count} line items have missing required fields",
                    ))

            if header.document_date is None:
                confidence -= 0.10

            confidence = max(round(confidence, 4), 0.0)

        except Exception as exc:
            failure_flags.append(FailureFlag(
                severity="ERROR",
                code="CORRUPTED_FILE",
                message=f"Failed to parse Word document: {type(exc).__name__}",
            ))
            confidence = 0.0

        return ParseResult(
            document_id=document_id,
            doc_type=DocType.INVOICE,
            parser_used=self.parser_name,
            parser_version="1.0.0",
            parse_confidence=confidence,
            header=header,
            line_items=line_items,
            failure_flags=failure_flags,
            raw_extracted_data=raw_data,
        )

    def _extract_header_info(
        self, paragraphs: list[str], failure_flags: list[FailureFlag]
    ) -> DocumentHeader:
        """Extract header information from document paragraphs."""
        header = DocumentHeader()
        full_text = "\n".join(paragraphs)

        # Vendor name
        vendor_match = _VENDOR_PATTERNS.search(full_text)
        if vendor_match:
            header.vendor_name = vendor_match.group(1).strip()
        else:
            failure_flags.append(FailureFlag(
                severity="WARNING",
                code="MISSING_VENDOR_NAME",
                message="Vendor name not found in document paragraphs",
            ))

        # Document number
        num_match = _DOC_NUMBER_PATTERNS.search(full_text)
        if num_match:
            header.document_number = num_match.group(1).strip()
        else:
            failure_flags.append(FailureFlag(
                severity="ERROR",
                code="MISSING_DOCUMENT_NUMBER",
                message="Document number could not be extracted",
            ))

        # Date
        date_match = _DATE_PATTERNS.search(full_text)
        if date_match:
            header.document_date = _parse_date_value(date_match.group(1))
        if header.document_date is None:
            failure_flags.append(FailureFlag(
                severity="ERROR",
                code="MISSING_DOCUMENT_DATE",
                message="Document date could not be extracted",
            ))

        # Total
        total_match = _TOTAL_PATTERNS.search(full_text)
        if total_match:
            header.total_amount = _parse_numeric(total_match.group(1))

        # GST
        gst_match = _GST_PATTERN.search(full_text)
        if gst_match:
            header.vendor_gst_id = gst_match.group(1)

        return header

    def _parse_table(self, table: Any, offset: int) -> list[LineItem]:
        """Parse a python-docx Table into LineItem objects."""
        rows = table.rows
        if not rows:
            return []

        # Convert all rows to lists of cell text
        all_rows = [[_cell_text(cell) for cell in row.cells] for row in rows]

        # Find header row
        header_idx = None
        for idx, row_cells in enumerate(all_rows):
            if _is_table_header_row(row_cells):
                header_idx = idx
                break

        if header_idx is None:
            header_idx = 0

        headers = all_rows[header_idx]
        col_mapping = _map_table_columns(headers)

        items: list[LineItem] = []
        for row_idx, row_cells in enumerate(all_rows[header_idx + 1:], start=1):
            # Skip empty rows
            if not any(c for c in row_cells):
                continue

            item = self._extract_line_item(row_cells, col_mapping, offset + row_idx)
            if item is not None:
                items.append(item)

        return items

    def _extract_line_item(
        self,
        cells: list[str],
        col_mapping: dict[str, int | None],
        line_number: int,
    ) -> LineItem | None:
        """Extract a single line item from table cells."""
        item_desc = None
        if col_mapping["item_desc"] is not None and col_mapping["item_desc"] < len(cells):
            val = cells[col_mapping["item_desc"]]
            if val:
                item_desc = val

        if not item_desc:
            return None

        quantity = None
        if col_mapping["quantity"] is not None and col_mapping["quantity"] < len(cells):
            quantity = _parse_numeric(cells[col_mapping["quantity"]])

        unit = None
        if col_mapping["unit"] is not None and col_mapping["unit"] < len(cells):
            val = cells[col_mapping["unit"]]
            if val:
                unit = val

        unit_price = None
        if col_mapping["unit_price"] is not None and col_mapping["unit_price"] < len(cells):
            unit_price = _parse_numeric(cells[col_mapping["unit_price"]])

        line_total = None
        if col_mapping["line_total"] is not None and col_mapping["line_total"] < len(cells):
            line_total = _parse_numeric(cells[col_mapping["line_total"]])

        if line_total is None and quantity is not None and unit_price is not None:
            line_total = quantity * unit_price

        field_confidences: dict[str, float] = {}
        if item_desc:
            field_confidences["item_desc"] = 1.0
        if quantity is not None:
            field_confidences["quantity"] = 1.0
        if unit_price is not None:
            field_confidences["unit_price"] = 1.0
        if line_total is not None:
            field_confidences["line_total"] = 1.0

        return LineItem(
            line_number=line_number,
            item_desc=item_desc,
            quantity=quantity,
            unit=unit,
            unit_price=unit_price,
            line_total=line_total,
            ordered_qty=None,
            received_qty=None,
            field_confidences=field_confidences,
        )
