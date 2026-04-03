"""Numeric extraction and comparison for Tool B."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from rapidfuzz import fuzz

NUMBER_PATTERN = re.compile(r"\b\d+(?:,\d{3})*(?:\.\d+)?\b")
CURRENCY_PATTERN = re.compile(r"[₹$€£]")


class NumericComparator:
    """Extract and compare numeric values across document versions."""

    def extract_numerics(self, file_path: Path) -> list[dict[str, Any]]:
        """Extract likely financial numeric values from supported documents."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(file_path)
        if suffix == ".docx":
            return self._extract_docx(file_path)
        if suffix in {".xlsx", ".xls", ".xlsm", ".csv"}:
            return self._extract_spreadsheet(file_path)
        return []

    def compare(
        self,
        current: list[dict[str, Any]],
        previous: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """Compare numeric values using context similarity."""
        if not previous:
            return []

        changes: list[dict[str, Any]] = []
        used_previous: set[int] = set()

        for current_item in current:
            best_idx: int | None = None
            best_score = -1.0
            for index, previous_item in enumerate(previous):
                if index in used_previous:
                    continue
                score = fuzz.token_sort_ratio(
                    str(current_item.get("context") or ""),
                    str(previous_item.get("context") or ""),
                )
                if score >= 80 and score > best_score:
                    best_idx = index
                    best_score = score

            if best_idx is None:
                continue

            previous_item = previous[best_idx]
            prev_value = float(previous_item["value"])
            current_value = float(current_item["value"])
            used_previous.add(best_idx)
            if prev_value == current_value:
                continue

            changes.append(
                {
                    "previous_value": prev_value,
                    "current_value": current_value,
                    "context": str(current_item.get("context") or previous_item.get("context") or "").strip(),
                    "change_pct": self._percentage_change(prev_value, current_value),
                }
            )

        return changes

    def _extract_pdf(self, file_path: Path) -> list[dict[str, Any]]:
        numerics: list[dict[str, Any]] = []
        try:
            with pdfplumber.open(str(file_path)) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    numerics.extend(
                        self._extract_from_text_block(
                            text=text,
                            location=f"Page {page_number}",
                        )
                    )
        except Exception:
            return []
        return numerics

    def _extract_docx(self, file_path: Path) -> list[dict[str, Any]]:
        numerics: list[dict[str, Any]] = []
        try:
            document = DocxDocument(str(file_path))
        except Exception:
            return []

        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            numerics.extend(
                self._extract_from_text_block(
                    text=paragraph.text or "",
                    location=f"Paragraph {paragraph_index}",
                )
            )

        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                row_text = " | ".join(cell.text or "" for cell in row.cells)
                numerics.extend(
                    self._extract_from_text_block(
                        text=row_text,
                        location=f"Table {table_index}, Row {row_index}",
                    )
                )
        return numerics

    def _extract_spreadsheet(self, file_path: Path) -> list[dict[str, Any]]:
        numerics: list[dict[str, Any]] = []
        try:
            sheets = pd.read_excel(file_path, sheet_name=None, header=None, dtype=object)
        except Exception:
            try:
                dataframe = pd.read_csv(file_path, header=None, dtype=object)
            except Exception:
                return []
            sheets = {"Sheet 1": dataframe}

        for sheet_name, frame in sheets.items():
            filled = frame.fillna("")
            for row_index, row in enumerate(filled.itertuples(index=False), start=1):
                cells = [str(cell).strip() for cell in row if str(cell).strip()]
                if not cells:
                    continue
                row_text = " | ".join(cells)
                numerics.extend(
                    self._extract_from_text_block(
                        text=row_text,
                        location=f"Sheet {sheet_name}, Row {row_index}",
                    )
                )
        return numerics

    def _extract_from_text_block(self, text: str, location: str) -> list[dict[str, Any]]:
        if not text or not text.strip():
            return []

        text = re.sub(r"\s+", " ", text).strip()
        matches: list[dict[str, Any]] = []
        for match in NUMBER_PATTERN.finditer(text):
            matched_text = match.group(0)
            before = text[max(0, match.start() - 20):match.start()]
            after = text[match.end():match.end() + 20]
            context = text[max(0, match.start() - 60):min(len(text), match.end() + 60)].strip()

            if self._is_percentage(text, match.end()):
                continue

            value = self._to_float(matched_text)
            if value is None:
                continue

            if self._is_year(value):
                continue

            if not self._looks_financial(value, matched_text, before, after):
                continue

            if self._looks_like_isolated_page_number(text, matched_text, value):
                continue

            matches.append(
                {
                    "value": value,
                    "context": context,
                    "location": location,
                }
            )
        return matches

    @staticmethod
    def _to_float(value: str) -> float | None:
        try:
            return float(value.replace(",", ""))
        except Exception:
            return None

    @staticmethod
    def _percentage_change(previous_value: float, current_value: float) -> float:
        if previous_value == 0:
            return 0.0 if current_value == 0 else 100.0
        return round(abs(((current_value - previous_value) / previous_value) * 100), 2)

    @staticmethod
    def _is_year(value: float) -> bool:
        return float(value).is_integer() and 1900 <= int(value) <= 2100

    @staticmethod
    def _is_percentage(text: str, end_index: int) -> bool:
        return end_index < len(text) and text[end_index] == "%"

    @staticmethod
    def _looks_financial(value: float, matched_text: str, before: str, after: str) -> bool:
        neighborhood = f"{before}{after}"
        has_currency = bool(CURRENCY_PATTERN.search(neighborhood))
        has_decimal = "." in matched_text
        return has_currency or has_decimal or value > 100

    @staticmethod
    def _looks_like_isolated_page_number(text: str, matched_text: str, value: float) -> bool:
        if not float(value).is_integer():
            return False
        if not 1 <= int(value) <= 999:
            return False
        return text.strip() == matched_text.strip()
