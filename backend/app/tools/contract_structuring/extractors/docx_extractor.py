"""
Word document extractor with table-first extraction and paragraph fallback.
"""

from __future__ import annotations

import logging
import re
from collections import defaultdict
from typing import List

from backend.app.tools.contract_structuring.extractors.base_extractor import (
    DocumentExtractionResult,
    RawTableResult,
)
from backend.app.tools.contract_structuring.extractors.table_normalizer import (
    LINE_ITEM_PATTERN,
    extract_currency_from_cell,
    normalize_tables_detailed,
)

logger = logging.getLogger(__name__)


def _clean_text(value) -> str:
    return str(value or "").strip()


def _sanitize_headers(values: list[str]) -> list[str]:
    seen: dict[str, int] = {}
    headers: list[str] = []
    for index, value in enumerate(values, start=1):
        header = _clean_text(value) or f"column_{index}"
        seen[header] = seen.get(header, 0) + 1
        if seen[header] > 1:
            header = f"{header}_{seen[header]}"
        headers.append(header)
    return headers


def _build_cell_text_cache(doc) -> dict[int, str]:
    cache: dict[int, str] = {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                key = id(cell._tc)
                text = _clean_text(cell.text)
                if text and key not in cache:
                    cache[key] = text
    return cache


def _read_row_cells(row, cell_text_cache: dict[int, str]) -> list[str]:
    values: list[str] = []
    for cell in row.cells:
        key = id(cell._tc)
        values.append(_clean_text(cell.text) or cell_text_cache.get(key, ""))
    return values


def _table_rows_from_doc_table(table, cell_text_cache: dict[int, str]) -> tuple[list[str], list[dict[str, str]]]:
    if not table.rows:
        return [], []

    header_values = _read_row_cells(table.rows[0], cell_text_cache)
    headers = _sanitize_headers(header_values)

    rows: list[dict[str, str]] = []
    for row in table.rows[1:]:
        row_values = _read_row_cells(row, cell_text_cache)
        row_dict = {
            headers[index]: row_values[index] if index < len(row_values) else ""
            for index in range(len(headers))
        }
        rows.append(row_dict)
    return headers, rows


def _paragraph_line_tables(doc) -> list[RawTableResult]:
    rows_by_page: dict[int, list[dict[str, str]]] = defaultdict(list)

    for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
        text = _clean_text(paragraph.text)
        if not text:
            continue

        segments = [segment.strip() for segment in re.split(r"[;\n]", text) if segment.strip()]
        if not segments:
            segments = [text]

        for segment in segments:
            for match in LINE_ITEM_PATTERN.finditer(segment):
                groups = match.groupdict()
                price_value, currency = extract_currency_from_cell(groups.get("price", ""))
                if price_value is None:
                    continue
                rows_by_page[((paragraph_index - 1) // 40) + 1].append(
                    {
                        "Item Description": groups.get("item", "").strip(),
                        "Quantity": groups.get("quantity", "").strip(),
                        "Unit": groups.get("unit", "").strip(),
                        "Unit Price": str(price_value),
                        "Currency": currency or "",
                    }
                )

    return [
        RawTableResult(
            source_page=page_number,
            extraction_method="DOCX_TABLE",
            raw_table_json=rows,
            table_confidence=0.45,
            column_count=len(rows[0]) if rows else 0,
            row_count=len(rows),
            source_name=f"docx_paragraphs_page_{page_number}",
            source_row_count=len(rows),
            failure_flags=["DOCX_PARAGRAPH_FALLBACK"],
        )
        for page_number, rows in sorted(rows_by_page.items())
        if rows
    ]


def extract_tables_from_docx(document_path: str) -> List[RawTableResult]:
    results: list[RawTableResult] = []
    try:
        from docx import Document

        doc = Document(document_path)
        cell_text_cache = _build_cell_text_cache(doc)
        for table_index, table in enumerate(doc.tables, start=1):
            headers, rows = _table_rows_from_doc_table(table, cell_text_cache)
            results.append(
                RawTableResult(
                    source_page=table_index,
                    extraction_method="DOCX_TABLE",
                    raw_table_json=rows,
                    table_confidence=0.70 if headers else 0.45,
                    column_count=len(headers),
                    row_count=len(rows),
                    source_name=f"table_{table_index}",
                    source_row_count=len(rows),
                )
            )
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
    return results


def extract_text_from_docx(document_path: str) -> str:
    try:
        from docx import Document

        doc = Document(document_path)
        fragments = [_clean_text(paragraph.text) for paragraph in doc.paragraphs if _clean_text(paragraph.text)]
        cell_text_cache = _build_cell_text_cache(doc)
        for table in doc.tables:
            for row in table.rows:
                values = [value for value in _read_row_cells(row, cell_text_cache) if value]
                if values:
                    fragments.append(" | ".join(values))
        return "\n".join(fragments)
    except Exception as exc:
        logger.error("DOCX text extraction failed: %s", exc)
        return ""


def _extract_docx_document(document_path: str) -> DocumentExtractionResult:
    from docx import Document

    raw_tables = extract_tables_from_docx(document_path)
    full_text = extract_text_from_docx(document_path)
    normalization = normalize_tables_detailed(raw_tables, stitched=False, document_text=full_text)
    tables = list(raw_tables)
    failure_flags = list(normalization.failure_flags)

    if not normalization.line_items:
        try:
            paragraph_tables = _paragraph_line_tables(Document(document_path))
            if paragraph_tables:
                tables.extend(paragraph_tables)
                normalization = normalize_tables_detailed(tables, stitched=False, document_text=full_text)
                failure_flags = list(normalization.failure_flags)
        except Exception as exc:
            logger.debug("DOCX paragraph fallback failed: %s", exc)

    return DocumentExtractionResult(
        tables=tables,
        line_items=normalization.line_items,
        clauses=[],
        confidence=normalization.confidence,
        failure_flags=failure_flags,
        text=full_text,
    )


class DocxExtractor:
    def extract_tables(self, document_path):
        return extract_tables_from_docx(str(document_path))

    def extract_text(self, document_path):
        return extract_text_from_docx(str(document_path))

    def extract(self, document_path):
        return _extract_docx_document(str(document_path))


DOCXExtractor = DocxExtractor
