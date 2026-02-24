from docx import Document
import sys
import os

filepath = sys.argv[1]
doc = Document(filepath)
for p in doc.paragraphs:
    print(p.text)

# Also extract tables
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
    print()
